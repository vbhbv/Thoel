import os
import re
import html
import time
import logging
import asyncio
import contextlib
from collections import deque
from pathlib import Path
from typing import Optional, List, Dict, Any, Deque

import asyncpg
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.constants import ParseMode
from telegram.request import HTTPXRequest
from telegram.error import TelegramError, Forbidden, RetryAfter, BadRequest
from telegram.ext import (
    Application,
    CommandHandler,
    CallbackQueryHandler,
    MessageHandler,
    ContextTypes,
    filters,
)

from pdf_epub_converter import (
    convert_epub_to_pdf,
    is_calibre_available,
    EbookConversionError,
)

# استيراد وحدات المعالجة الخارجية في أعلى الملف (بدل منتصفه كما كان سابقًا)
import files_handler
import audio_handler

# ----------------------------------------------------------------------
# الإعدادات العامة ومتغيرات البيئة
# ----------------------------------------------------------------------

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

BOT_TOKEN = os.environ.get("BOT_TOKEN")
if not BOT_TOKEN:
    raise RuntimeError("يجب ضبط متغير البيئة BOT_TOKEN")

DATABASE_URL = os.environ.get("DATABASE_URL")
if not DATABASE_URL:
    raise RuntimeError("يجب ضبط متغير البيئة DATABASE_URL المربوط بقاعدة PostgreSQL")


def _parse_admin_ids() -> List[int]:
    """رفض افتراضي (قائمة فارغة) إن لم يُضبط ADMIN_IDS، بدل معرف وهمي
    ("0") يوهم بوجود حماية غير حقيقية."""
    raw = os.environ.get("ADMIN_IDS", "").strip()
    if not raw:
        logger.critical(
            "⚠️ متغير البيئة ADMIN_IDS غير مضبوط! لوحة تحكم الأدمن ستكون "
            "معطّلة بالكامل حتى تضبط هذا المتغير بمعرفك الرقمي."
        )
        return []
    ids: List[int] = []
    for part in raw.split(","):
        part = part.strip()
        if not part:
            continue
        try:
            ids.append(int(part))
        except ValueError:
            logger.critical(f"⚠️ قيمة غير صالحة في ADMIN_IDS تم تجاهلها: '{part}'")
    return ids


ADMIN_IDS: List[int] = _parse_admin_ids()

BASE_DIR = Path(__file__).parent
DOWNLOADS_DIR = BASE_DIR / "downloads"
CONVERTED_DIR = BASE_DIR / "converted"
DOWNLOADS_DIR.mkdir(exist_ok=True)
CONVERTED_DIR.mkdir(exist_ok=True)

FILE_MAX_AGE = 60 * 60  # ساعة واحدة
CLEANUP_INTERVAL = 30 * 60  # 30 دقيقة
MAX_FILE_SIZE_MB = 20
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
MAX_MERGE_FILES = int(os.environ.get("MAX_MERGE_FILES", "20"))

AUDIO_FORMATS = ["mp3", "m4a", "flac", "ogg", "wav", "aac"]
DOC_EXTENSIONS = {".doc", ".docx"}
PDF_EXTENSION = ".pdf"
EPUB_EXTENSION = ".epub"
AUDIO_EXTENSIONS = {".mp3", ".wav", ".ogg", ".flac", ".m4a", ".aac", ".opus", ".wma"}
VIDEO_EXTENSIONS = {".mp4", ".mkv", ".avi", ".mov", ".flv", ".wmv", ".3gp", ".webm"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tiff", ".tif"}

DB_POOL_MIN_SIZE = int(os.environ.get("DB_POOL_MIN_SIZE", "2"))
DB_POOL_MAX_SIZE = int(os.environ.get("DB_POOL_MAX_SIZE", "10"))
DB_COMMAND_TIMEOUT = float(os.environ.get("DB_COMMAND_TIMEOUT", "30"))
DB_CONNECT_RETRIES = int(os.environ.get("DB_CONNECT_RETRIES", "5"))
DB_CONNECT_RETRY_DELAY = float(os.environ.get("DB_CONNECT_RETRY_DELAY", "3"))

RATE_LIMIT_MAX_REQUESTS = int(os.environ.get("RATE_LIMIT_MAX_REQUESTS", "8"))
RATE_LIMIT_WINDOW_SECONDS = float(os.environ.get("RATE_LIMIT_WINDOW_SECONDS", "20"))

BAN_CACHE_TTL = float(os.environ.get("BAN_CACHE_TTL_SECONDS", "30"))
SETTINGS_CACHE_TTL = float(os.environ.get("SETTINGS_CACHE_TTL_SECONDS", "60"))

BROADCAST_CONCURRENCY = int(os.environ.get("BROADCAST_CONCURRENCY", "15"))
BROADCAST_DELAY_PER_MESSAGE = float(os.environ.get("BROADCAST_DELAY_PER_MESSAGE", "0.05"))

# مهلة قصوى (بالثواني) لأي أمر خارجي (ffmpeg/libreoffice)، لمنع تعليق
# العملية إلى الأبد على ملف تالف أو ضخم بشكل غير طبيعي.
CMD_TIMEOUT_SECONDS = float(os.environ.get("CMD_TIMEOUT_SECONDS", "300"))

# حد أقصى لعدد العمليات "الثقيلة" (ضغط PDF، تحويل فيديو، TTS، دمج وسائط)
# العاملة بالتوازي، لحماية موارد CPU/RAM من الاستنزاف الكامل تحت الضغط.
HEAVY_TASK_CONCURRENCY = int(os.environ.get("HEAVY_TASK_CONCURRENCY", "3"))

MAX_IMAGES_PER_BATCH = int(os.environ.get("MAX_IMAGES_PER_BATCH", "30"))

# حماية من "قنابل الضغط" (صور بأبعاد ضخمة جدًا تستنزف الذاكرة عند فك
# ضغطها بواسطة PIL، حتى لو كان حجم الملف نفسه صغيرًا على القرص).
MAX_IMAGE_PIXELS = int(os.environ.get("MAX_IMAGE_PIXELS", str(40_000_000)))  # ~40 ميغابكسل
from PIL import Image as _PILImageConfig  # noqa: E402
_PILImageConfig.MAX_IMAGE_PIXELS = MAX_IMAGE_PIXELS


class HeavyTaskGuard:
    """مدير سياق يحدّ من عدد العمليات الثقيلة العاملة بالتوازي (استخدام:
    `async with HeavyTaskGuard(): ...`) لمنع استنزاف موارد الخادم بالكامل
    عند تلقي عدة طلبات ثقيلة من مستخدمين مختلفين في نفس اللحظة."""

    _semaphore = asyncio.Semaphore(HEAVY_TASK_CONCURRENCY)

    async def __aenter__(self):
        await self._semaphore.acquire()
        return self

    async def __aexit__(self, exc_type, exc, tb):
        self._semaphore.release()


# ----------------------------------------------------------------------
# أداة كاش بسيطة بصلاحية زمنية (TTL) - بدون أي اعتمادية خارجية
# ----------------------------------------------------------------------

_CACHE_NONE_SENTINEL = object()


class TTLCache:
    """كاش في الذاكرة مع صلاحية زمنية. يعمل ضمن عملية واحدة (in-memory)،
    كافٍ لخدمة تعمل كنسخة واحدة. عند التوسع لعدة نسخ متزامنة يُفضّل
    استبداله بـ Redis لمشاركة الكاش بين النسخ."""

    def __init__(self, ttl_seconds: float):
        self._ttl = ttl_seconds
        self._store: Dict[Any, tuple] = {}

    def get(self, key: Any) -> Any:
        entry = self._store.get(key)
        if entry is None:
            return None
        value, expires_at = entry
        if time.monotonic() > expires_at:
            self._store.pop(key, None)
            return None
        return value

    def set(self, key: Any, value: Any) -> None:
        self._store[key] = (value, time.monotonic() + self._ttl)

    def invalidate(self, key: Optional[Any] = None) -> None:
        if key is None:
            self._store.clear()
        else:
            self._store.pop(key, None)


_ban_cache = TTLCache(ttl_seconds=BAN_CACHE_TTL)
_settings_cache = TTLCache(ttl_seconds=SETTINGS_CACHE_TTL)


# ----------------------------------------------------------------------
# محدد المعدل (Rate Limiter) لحماية البوت من إساءة الاستخدام
# ----------------------------------------------------------------------

class RateLimiter:
    """محدد معدل بسيط (Sliding Window) في الذاكرة."""

    def __init__(self, max_requests: int, window_seconds: float):
        self.max_requests = max_requests
        self.window_seconds = window_seconds
        self._hits: Dict[int, Deque[float]] = {}

    def allow(self, user_id: int) -> bool:
        now = time.monotonic()
        window_start = now - self.window_seconds

        hits = self._hits.get(user_id)
        if hits is None:
            hits = deque()
            self._hits[user_id] = hits

        while hits and hits[0] < window_start:
            hits.popleft()

        if len(hits) >= self.max_requests:
            return False

        hits.append(now)
        return True


_rate_limiter = RateLimiter(RATE_LIMIT_MAX_REQUESTS, RATE_LIMIT_WINDOW_SECONDS)


def is_rate_limited(user_id: int) -> bool:
    return not _rate_limiter.allow(user_id)


def prune_rate_limiter() -> int:
    """يزيل سجلات المستخدمين غير النشطين من محدد المعدل لتفادي أي تراكم
    في الذاكرة على المدى الطويل جدًا. تُستدعى من مهمة التنظيف الدورية."""
    now = time.monotonic()
    window_start = now - _rate_limiter.window_seconds
    stale_keys = [
        uid for uid, hits in _rate_limiter._hits.items()
        if not hits or hits[-1] < window_start
    ]
    for uid in stale_keys:
        del _rate_limiter._hits[uid]
    return len(stale_keys)


# ----------------------------------------------------------------------
# تهيئة وإدارة اتصال قاعدة البيانات (Connection Pool)
# ----------------------------------------------------------------------

_pool: Optional[asyncpg.Pool] = None


def _get_pool() -> Optional[asyncpg.Pool]:
    if _pool is None:
        logger.debug("⚠️ محاولة استخدام قاعدة البيانات قبل تهيئتها أو أثناء عدم توفرها.")
    return _pool


def fix_database_url(url: str) -> str:
    if url.startswith("postgres://"):
        return url.replace("postgres://", "postgresql://", 1)
    return url


_COMMAND_TAG_COUNT_RE = re.compile(r"(\d+)\s*$")


def _parse_command_tag_count(tag: str) -> int:
    match = _COMMAND_TAG_COUNT_RE.search(tag or "")
    return int(match.group(1)) if match else 0


async def init_db() -> bool:
    """ينشئ Pool اتصالات بقاعدة البيانات (مع إعادة محاولة تلقائية) ويهيئ
    الجداول والفهارس. يعيد True عند النجاح، False عند الفشل النهائي."""
    global _pool

    url = fix_database_url(DATABASE_URL)

    last_error: Optional[Exception] = None
    for attempt in range(1, DB_CONNECT_RETRIES + 1):
        try:
            _pool = await asyncpg.create_pool(
                dsn=url,
                min_size=DB_POOL_MIN_SIZE,
                max_size=DB_POOL_MAX_SIZE,
                command_timeout=DB_COMMAND_TIMEOUT,
            )
            break
        except Exception as e:
            last_error = e
            logger.warning(
                f"⏳ محاولة الاتصال بقاعدة البيانات {attempt}/{DB_CONNECT_RETRIES} فشلت، "
                f"إعادة المحاولة خلال {DB_CONNECT_RETRY_DELAY} ثانية..."
            )
            await asyncio.sleep(DB_CONNECT_RETRY_DELAY)
    else:
        logger.critical(f"❌ تعذر الاتصال بقاعدة البيانات بعد {DB_CONNECT_RETRIES} محاولات: {last_error}")
        return False

    try:
        async with _pool.acquire() as conn:
            await conn.execute("""
                CREATE TABLE IF NOT EXISTS users (
                    user_id BIGINT PRIMARY KEY,
                    username TEXT,
                    joined_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    last_active TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    is_banned BOOLEAN DEFAULT FALSE,
                    blocked_bot BOOLEAN DEFAULT FALSE
                );
            """)
            await conn.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS last_active TIMESTAMP DEFAULT CURRENT_TIMESTAMP;")
            await conn.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS blocked_bot BOOLEAN DEFAULT FALSE;")

            await conn.execute("""
                CREATE TABLE IF NOT EXISTS stats_log (
                    id SERIAL PRIMARY KEY,
                    action_type TEXT NOT NULL,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
            """)

            await conn.execute("""
                CREATE TABLE IF NOT EXISTS settings (
                    key TEXT PRIMARY KEY,
                    value TEXT
                );
            """)

            await conn.execute("CREATE INDEX IF NOT EXISTS idx_stats_log_timestamp ON stats_log(timestamp);")
            await conn.execute("CREATE INDEX IF NOT EXISTS idx_users_banned ON users(is_banned) WHERE is_banned = TRUE;")

        logger.info(
            f"✅ تم الاتصال بقاعدة البيانات (Pool: {DB_POOL_MIN_SIZE}-{DB_POOL_MAX_SIZE}) وتهيئة الجداول بنجاح."
        )
        return True
    except Exception as e:
        logger.critical(f"❌ فشل تهيئة جداول قاعدة البيانات: {e}")
        return False


async def close_db() -> None:
    global _pool
    if _pool is not None:
        await _pool.close()
        _pool = None
        logger.info("🔒 تم إغلاق اتصال قاعدة البيانات بأمان.")


# ----------------------------------------------------------------------
# دالات إدارة المستخدمين
# ----------------------------------------------------------------------

async def register_user(user_id: int, username: Optional[str]) -> None:
    pool = _get_pool()
    if pool is None:
        return
    try:
        async with pool.acquire() as conn:
            await conn.execute(
                """
                INSERT INTO users (user_id, username, last_active)
                VALUES ($1, $2, CURRENT_TIMESTAMP)
                ON CONFLICT (user_id) DO UPDATE
                SET username = EXCLUDED.username, last_active = CURRENT_TIMESTAMP;
                """,
                user_id, username,
            )
    except Exception:
        logger.exception(f"فشل تسجيل المستخدم {user_id}")


async def is_user_banned(user_id: int) -> bool:
    cached = _ban_cache.get(user_id)
    if cached is not None:
        return cached

    pool = _get_pool()
    if pool is None:
        return False

    try:
        async with pool.acquire() as conn:
            row = await conn.fetchrow("SELECT is_banned FROM users WHERE user_id = $1;", user_id)
        result = bool(row["is_banned"]) if row else False
    except Exception:
        logger.exception(f"خطأ أثناء فحص حظر المستخدم {user_id}")
        result = False  # فشل مفتوح لتفادي تعطل البوت بالكامل بسبب عطل مؤقت بالقاعدة

    _ban_cache.set(user_id, result)
    return result


async def set_user_ban_status(user_id: int, banned: bool) -> None:
    pool = _get_pool()
    if pool is None:
        raise RuntimeError("قاعدة البيانات غير متاحة حاليًا.")
    async with pool.acquire() as conn:
        await conn.execute("UPDATE users SET is_banned = $2 WHERE user_id = $1;", user_id, banned)
    _ban_cache.invalidate(user_id)


async def log_action(action_type: str) -> None:
    pool = _get_pool()
    if pool is None:
        return
    try:
        async with pool.acquire() as conn:
            await conn.execute("INSERT INTO stats_log (action_type) VALUES ($1);", action_type)
    except Exception:
        logger.exception(f"خطأ أثناء تسجيل العملية {action_type}")


async def get_setting(key: str) -> Optional[str]:
    cached = _settings_cache.get(key)
    if cached is not None:
        return None if cached is _CACHE_NONE_SENTINEL else cached

    pool = _get_pool()
    if pool is None:
        return None

    try:
        async with pool.acquire() as conn:
            val = await conn.fetchval("SELECT value FROM settings WHERE key = $1;", key)
    except Exception:
        logger.exception(f"خطأ أثناء جلب الإعداد {key}")
        return None

    _settings_cache.set(key, val if val is not None else _CACHE_NONE_SENTINEL)
    return val


async def set_setting(key: str, value: str) -> None:
    pool = _get_pool()
    if pool is None:
        raise RuntimeError("قاعدة البيانات غير متاحة حاليًا.")
    async with pool.acquire() as conn:
        await conn.execute(
            "INSERT INTO settings (key, value) VALUES ($1, $2) ON CONFLICT (key) DO UPDATE SET value = $2;",
            key, value,
        )
    _settings_cache.invalidate(key)


# ----------------------------------------------------------------------
# فحص الاشتراك الإجباري الاحترافي
# ----------------------------------------------------------------------

async def check_force_subscription(user_id: int, channel_username: str, context: ContextTypes.DEFAULT_TYPE) -> bool:
    """يتحقق من عضوية المستخدم الفعلية في القناة المحددة عبر Telegram API."""
    clean_username = channel_username.replace("@", "").strip()
    chat_id = f"@{clean_username}"
    try:
        member = await context.bot.get_chat_member(chat_id=chat_id, user_id=user_id)
        return member.status in ("member", "administrator", "creator")
    except Exception as e:
        logger.warning(f"تعذر فحص الاشتراك الإجباري في {chat_id} (يُحتمل أن البوت ليس مشرفًا فيها): {e}")
        # فشل مفتوح: لا نمنع كل المستخدمين من استخدام البوت بسبب إعداد قناة خاطئ
        return True


async def _is_subscription_required_and_unmet(user_id: int, context: ContextTypes.DEFAULT_TYPE) -> Optional[str]:
    """يرجع يوزر القناة إن كان الاشتراك مفروضًا وغير مُستوفى، أو None إن
    كان كل شيء سليمًا (الأدمن مستثنى دائمًا)."""
    if user_id in ADMIN_IDS:
        return None
    channel_username = await get_setting("force_channel")
    if not channel_username:
        return None
    if await check_force_subscription(user_id, channel_username, context):
        return None
    return channel_username


def get_sub_keyboard(channel_username: str) -> InlineKeyboardMarkup:
    clean_username = channel_username.replace("@", "").strip()
    buttons = [
        [InlineKeyboardButton("📢 اشترك في القناة أولاً", url=f"https://t.me/{clean_username}")],
        [InlineKeyboardButton("✅ تم الاشتراك (تفعيل)", callback_data="check_sub_again")],
    ]
    return InlineKeyboardMarkup(buttons)


# ----------------------------------------------------------------------
# دالات حماية موحدة تُستخدم في بداية كل معالج
# ----------------------------------------------------------------------

async def _reject_if_banned(update: Update) -> bool:
    """يعيد True إن كان المستخدم محظورًا (ويرسل له رسالة موحّدة)."""
    user_id = update.effective_user.id
    if await is_user_banned(user_id):
        if update.effective_message:
            await update.effective_message.reply_text("🚫 عذرًا، حسابك محظور من استخدام هذا البوت.")
        return True
    return False


async def _reject_if_rate_limited(update: Update) -> bool:
    user_id = update.effective_user.id
    if is_rate_limited(user_id):
        if update.effective_message:
            await update.effective_message.reply_text(
                "⏳ تُرسل الطلبات بسرعة كبيرة. الرجاء الانتظار قليلًا ثم إعادة المحاولة."
            )
        return True
    return False


# ----------------------------------------------------------------------
# أدوات مساعدة للتشغيل والتحويل الفني والوسائط
# ----------------------------------------------------------------------

async def run_cmd(*args: str, timeout: float = CMD_TIMEOUT_SECONDS) -> tuple:
    """تشغيل أمر خارجي مع مهلة زمنية قصوى، لمنع تعليق العملية إلى الأبد
    على ملف تالف أو معقد بشكل غير طبيعي (وبالتالي استنزاف موارد الخادم)."""
    process = await asyncio.create_subprocess_exec(
        *args,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    try:
        stdout, stderr = await asyncio.wait_for(process.communicate(), timeout=timeout)
    except asyncio.TimeoutError:
        process.kill()
        await process.wait()
        return 1, "", f"انتهت المهلة الزمنية القصوى ({timeout:.0f} ثانية) لتنفيذ: {args[0]}"
    return process.returncode, stdout.decode(errors="ignore"), stderr.decode(errors="ignore")


async def download_telegram_file(context: ContextTypes.DEFAULT_TYPE, file_id: str, file_unique_id: str, filename: str) -> Path:
    file_obj = await context.bot.get_file(file_id)
    ext = Path(filename).suffix or ""
    local_path = DOWNLOADS_DIR / f"{file_unique_id}{ext}"
    await file_obj.download_to_drive(custom_path=str(local_path))
    return local_path


async def _check_size_or_reject(update: Update, file_size: Optional[int]) -> bool:
    """يتحقق من حجم الملف قبل التحميل. يعيد True إن كان الحجم مقبولاً."""
    if file_size and file_size > MAX_FILE_SIZE_BYTES:
        await update.message.reply_text(f"⚠️ الملف أكبر من الحد المسموح ({MAX_FILE_SIZE_MB} ميغابايت).")
        return False
    return True


def apply_audio_metadata(audio_path: Path, title: str = None, artist: str = None, album_art_path: Path = None):
    from PIL import Image
    from io import BytesIO

    ext = audio_path.suffix.lower()
    image_bytes = None

    if album_art_path and album_art_path.exists():
        try:
            with Image.open(album_art_path) as img:
                if img.mode in ("RGBA", "P", "LA") or img.format == "WEBP":
                    img = img.convert("RGB")
                buffer = BytesIO()
                img.save(buffer, format="JPEG", quality=95)
                image_bytes = buffer.getvalue()
        except Exception as img_err:
            logger.error(f"خطأ أثناء معالجة الصورة وتحويلها إلى JPEG: {img_err}")

    try:
        if ext == ".mp3":
            from mutagen.id3 import ID3, TIT2, TPE1, APIC, ID3NoHeaderError
            try:
                audio = ID3(str(audio_path))
            except ID3NoHeaderError:
                audio = ID3()
            if title:
                audio.add(TIT2(encoding=3, text=title))
            if artist:
                audio.add(TPE1(encoding=3, text=artist))
            if image_bytes:
                audio.delall("APIC")
                audio.add(APIC(encoding=3, mime="image/jpeg", type=3, desc="Cover", data=image_bytes))
            audio.save(str(audio_path))

        elif ext in [".m4a", ".aac"]:
            from mutagen.mp4 import MP4, MP4Cover
            try:
                audio = MP4(str(audio_path))
                if title:
                    audio["\xa9nam"] = title
                if artist:
                    audio["\xa9ART"] = artist
                if image_bytes:
                    audio["covr"] = [MP4Cover(image_bytes, imageformat=MP4Cover.FORMAT_JPEG)]
                audio.save()
            except Exception as mp4_err:
                logger.error(f"فشلت معالجة حاوية MP4/M4A: {mp4_err}")

        elif ext == ".flac":
            from mutagen.flac import FLAC, Picture
            audio = FLAC(str(audio_path))
            if title:
                audio["title"] = title
            if artist:
                audio["artist"] = artist
            if image_bytes:
                pic = Picture()
                pic.data = image_bytes
                pic.type = 3
                pic.mime = "image/jpeg"
                pic.description = "Cover"
                audio.clear_pictures()
                audio.add_picture(pic)
            audio.save()

        elif ext in [".ogg", ".opus"]:
            from mutagen.oggvorbis import OggVorbis
            import base64
            from mutagen.flac import Picture
            audio = OggVorbis(str(audio_path))
            if title:
                audio["title"] = title
            if artist:
                audio["artist"] = artist
            if image_bytes:
                pic = Picture()
                pic.data = image_bytes
                pic.type = 3
                pic.mime = "image/jpeg"
                pic.description = "Cover"
                audio["metadata_block_picture"] = [base64.b64encode(pic.write()).decode("ascii")]
            audio.save()
    except Exception as e:
        logger.error(f"خطأ غير متوقع في Mutagen للصيغة {ext}: {e}")


async def convert_docx_to_pdf(input_path: Path, out_dir: Path) -> Path:
    code, out, err = await run_cmd(
        "libreoffice", "--headless", "--norestore",
        "--convert-to", "pdf", "--outdir", str(out_dir), str(input_path)
    )
    result = out_dir / (input_path.stem + ".pdf")
    if code != 0 or not result.exists():
        raise RuntimeError(f"فشل تحويل Word إلى PDF: {err or out}")
    return result


async def convert_pdf_to_docx(input_path: Path, out_dir: Path) -> Path:
    output_path = out_dir / (input_path.stem + ".docx")

    def _convert():
        from pdf2docx import Converter
        cv = Converter(str(input_path))
        cv.convert(str(output_path))
        cv.close()

    await asyncio.get_running_loop().run_in_executor(None, _convert)
    if not output_path.exists():
        raise RuntimeError("فشل تحويل PDF إلى Word")
    return output_path


async def convert_audio(input_path: Path, out_dir: Path, target_format: str) -> Path:
    output_path = out_dir / (input_path.stem + f".{target_format}")
    code, out, err = await run_cmd(
        "ffmpeg", "-y", "-i", str(input_path),
        "-vn", "-ar", "44100", "-ac", "2",
        str(output_path)
    )
    if code != 0 or not output_path.exists():
        raise RuntimeError(f"فشل تحويل الصوت: {err[-500:]}")
    return output_path


async def convert_video_to_audio(input_path: Path, out_dir: Path) -> Path:
    output_path = out_dir / (input_path.stem + ".mp3")
    code, out, err = await run_cmd(
        "ffmpeg", "-y", "-i", str(input_path),
        "-vn", "-acodec", "libmp3lame", "-q:a", "2",
        "-ar", "44100", "-ac", "2",
        str(output_path)
    )
    if code != 0 or not output_path.exists():
        raise RuntimeError(f"فشل استخراج الصوت من الفيديو: {err[-500:]}")
    return output_path


async def convert_images_to_pdf(input_paths: List[Path], out_dir: Path, base_name: str) -> Path:
    if len(input_paths) > MAX_IMAGES_PER_BATCH:
        raise RuntimeError(f"عدد الصور يتجاوز الحد الأقصى المسموح ({MAX_IMAGES_PER_BATCH} صورة دفعة واحدة).")

    output_path = out_dir / (base_name + ".pdf")

    def _convert():
        from PIL import Image, UnidentifiedImageError
        import img2pdf
        processed_paths = []
        try:
            for path in input_paths:
                try:
                    with Image.open(path) as img:
                        if img.mode in ("RGBA", "P", "LA"):
                            img = img.convert("RGB")
                        tmp_path = path.with_suffix(".conv.jpg")
                        img.save(tmp_path, "JPEG", quality=95)
                        processed_paths.append(str(tmp_path))
                except Image.DecompressionBombError:
                    raise RuntimeError(f"تم رفض الصورة {path.name}: أبعادها كبيرة جدًا (حماية من استنزاف الذاكرة).")
                except UnidentifiedImageError:
                    raise RuntimeError(f"تعذر قراءة الصورة {path.name}: الملف تالف أو ليس صورة صالحة.")

            pdf_bytes = img2pdf.convert(processed_paths)
            with open(output_path, "wb") as f:
                f.write(pdf_bytes)
        finally:
            for p in processed_paths:
                Path(p).unlink(missing_ok=True)

    await asyncio.get_running_loop().run_in_executor(None, _convert)
    return output_path


async def convert_images_to_docx(input_paths: List[Path], out_dir: Path, base_name: str) -> Path:
    if len(input_paths) > MAX_IMAGES_PER_BATCH:
        raise RuntimeError(f"عدد الصور يتجاوز الحد الأقصى المسموح ({MAX_IMAGES_PER_BATCH} صورة دفعة واحدة).")

    output_path = out_dir / (base_name + ".docx")

    def _convert():
        from docx import Document
        from docx.shared import Inches
        from PIL import Image, UnidentifiedImageError
        doc = Document()
        for path in input_paths:
            try:
                with Image.open(path) as img:
                    if img.mode in ("RGBA", "P", "LA"):
                        img = img.convert("RGB")
                        fixed_path = path.with_suffix(".conv.jpg")
                        img.save(fixed_path, "JPEG", quality=95)
                        source = fixed_path
                    else:
                        source = path
            except Image.DecompressionBombError:
                raise RuntimeError(f"تم رفض الصورة {path.name}: أبعادها كبيرة جدًا (حماية من استنزاف الذاكرة).")
            except UnidentifiedImageError:
                raise RuntimeError(f"تعذر قراءة الصورة {path.name}: الملف تالف أو ليس صورة صالحة.")

            doc.add_picture(str(source), width=Inches(6))
            if source != path:
                source.unlink(missing_ok=True)
        doc.save(str(output_path))

    await asyncio.get_running_loop().run_in_executor(None, _convert)
    return output_path


def encrypt_pdf_file(input_path: Path, output_path: Path, password: str):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(str(input_path))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    with open(output_path, "wb") as f:
        writer.write(f)


def make_progress_bar(percent: int) -> str:
    total_blocks = 10
    filled_blocks = int(percent / 10)
    empty_blocks = total_blocks - filled_blocks
    bar = "█" * filled_blocks + "░" * empty_blocks
    return f"{bar} {percent}%"


async def compress_pdf_file_async(input_path: Path, output_path: Path, status_msg) -> None:
    """يضغط ملف PDF صفحة بصفحة. مُنفَّذة بالكامل داخل executor لتفادي
    تجميد حلقة الأحداث (event loop) أمام كل المستخدمين الآخرين أثناء
    معالجة ملف PDF كبير - هذا كان أخطر عيب في النسخة الأصلية."""
    from pypdf import PdfReader, PdfWriter

    loop = asyncio.get_running_loop()
    last_update_time = {"t": 0.0}

    def _report_progress(idx: int, total: int) -> None:
        now = time.time()
        if now - last_update_time["t"] < 1.5 and idx != total:
            return
        last_update_time["t"] = now
        percent = int((idx / total) * 100) if total else 100
        bar = make_progress_bar(percent)

        async def _edit():
            with contextlib.suppress(Exception):
                await status_msg.edit_text(
                    f"⏳ <b>جاري معالجة وضغط صفحات الـ PDF...</b>\n\n"
                    f"📄 الصفحة: <code>{idx}</code> من <code>{total}</code>\n<code>{bar}</code>",
                    parse_mode=ParseMode.HTML,
                )

        asyncio.run_coroutine_threadsafe(_edit(), loop)

    def _compress_sync():
        reader = PdfReader(str(input_path))
        writer = PdfWriter()
        total_pages = len(reader.pages)
        for idx, page in enumerate(reader.pages, start=1):
            try:
                page.compress_content_streams()
            except Exception as page_err:
                logger.warning(f"تم تخطي ضغط الصفحة {idx} لتجنب الانهيار: {page_err}")
            writer.add_page(page)
            _report_progress(idx, total_pages)
        with open(output_path, "wb") as f:
            writer.write(f)

    await loop.run_in_executor(None, _compress_sync)


async def trim_audio_file(input_path: Path, output_path: Path, start_time: str, end_time: str):
    """قص مقطع صوتي. نعيد الترميز (Output Seeking بعد -i) بدل النسخ
    المباشر (-acodec copy)، لأن النسخ المباشر يقص فقط عند أقرب إطار
    مفتاح (keyframe) وقد يُنتج قصًا غير دقيق بثوانٍ كاملة أحيانًا حسب
    الصيغة؛ إعادة الترميز أبطأ قليلًا لكنها دقيقة عند نقطة القص تمامًا."""
    code, out, err = await run_cmd(
        "ffmpeg", "-y", "-i", str(input_path),
        "-ss", start_time, "-to", end_time,
        str(output_path),
    )
    if code != 0 or not output_path.exists():
        raise RuntimeError(f"فشل قص الصوت، تأكد من صحة كتابة الوقت المتطابق: {err[-300:]}")


# ----------------------------------------------------------------------
# فلاتر ومعالجات الصوت (سرعة، دمج، تحكّم بالصوت)
# ----------------------------------------------------------------------

async def change_audio_speed(input_path: Path, output_path: Path, speed: float):
    code, out, err = await run_cmd(
        "ffmpeg", "-y", "-i", str(input_path),
        "-filter:a", f"atempo={speed}",
        "-vn", str(output_path)
    )
    if code != 0 or not output_path.exists():
        raise RuntimeError(f"فشل تغيير سرعة الصوت: {err[-300:]}")


async def merge_audio_files(input_paths: List[Path], output_path: Path):
    """دمج ملفات صوتية متعددة. نُطبّع كل الملفات أولًا إلى صيغة/معدل
    عينات موحدين (WAV PCM) قبل الدمج، لأن فلتر concat في ffmpeg يتطلب
    تطابق معاملات التدفقات؛ دمج صيغ مختلطة (mp3+wav+m4a) مباشرة عبره
    قد يفشل بصمت أو يُنتج ملفًا صوتيًا تالفًا/مشوّهًا."""
    if len(input_paths) < 2:
        raise RuntimeError("يجب توفير ملفين على الأقل للدمج.")

    normalized_paths: List[Path] = []
    try:
        for i, p in enumerate(input_paths):
            norm_path = output_path.parent / f"_norm_{i}_{p.stem}.wav"
            code, out, err = await run_cmd(
                "ffmpeg", "-y", "-i", str(p),
                "-ar", "44100", "-ac", "2", "-c:a", "pcm_s16le",
                str(norm_path),
            )
            if code != 0 or not norm_path.exists():
                raise RuntimeError(f"فشل تجهيز الملف {p.name} قبل الدمج: {err[-300:]}")
            normalized_paths.append(norm_path)

        inputs = []
        for p in normalized_paths:
            inputs.extend(["-i", str(p)])
        filter_complex = f"concat=n={len(normalized_paths)}:v=0:a=1[a]"
        code, out, err = await run_cmd(
            "ffmpeg", "-y", *inputs,
            "-filter_complex", filter_complex,
            "-map", "[a]", "-c:a", "libmp3lame", "-q:a", "2",
            str(output_path),
        )
        if code != 0 or not output_path.exists():
            raise RuntimeError(f"فشل دمج الملفات الصوتية: {err[-300:]}")
    finally:
        for p in normalized_paths:
            p.unlink(missing_ok=True)


async def change_audio_volume(input_path: Path, output_path: Path, volume_db: float):
    code, out, err = await run_cmd(
        "ffmpeg", "-y", "-i", str(input_path),
        "-filter:a", f"volume={volume_db}dB",
        str(output_path)
    )
    if code != 0 or not output_path.exists():
        raise RuntimeError(f"فشل تعديل مستوى الصوت: {err[-300:]}")


async def process_epub_to_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE, tg_file, filename: str):
    if not await _check_size_or_reject(update, getattr(tg_file, "file_size", None)):
        return
    msg = await update.message.reply_text("⏳ جاري تحويل الكتاب الإلكتروني...")
    try:
        if not is_calibre_available():
            await msg.edit_text("❌ برمجية Calibre غير متوفرة على البيئة السحابية حاليًا.")
            return
        lp = await download_telegram_file(context, tg_file.file_id, tg_file.file_unique_id, filename)
        res = await convert_epub_to_pdf(lp, CONVERTED_DIR)
        await log_action("epub_to_pdf")
        with open(res, "rb") as f:
            await update.message.reply_document(document=f, filename=res.name)
        await msg.delete()
    except EbookConversionError as e:
        await msg.edit_text(f"❌ {html.escape(str(e))}")
    except Exception as e:
        logger.exception("خطأ أثناء تحويل EPUB إلى PDF")
        await msg.edit_text(f"❌ خطأ: {html.escape(str(e))}")


async def finalize_and_send_audio(chat_id: int, context: ContextTypes.DEFAULT_TYPE):
    audio_path_str = context.user_data.get("ready_audio_path")
    if not audio_path_str:
        await context.bot.send_message(chat_id=chat_id, text="❌ لم يتم العثور على ملف جاهز للبث.")
        return

    audio_path = Path(audio_path_str)
    title = context.user_data.get("meta_title")
    artist = context.user_data.get("meta_artist")
    art_path_str = context.user_data.get("meta_art_path")
    art_path = Path(art_path_str) if art_path_str else None

    await asyncio.get_running_loop().run_in_executor(None, apply_audio_metadata, audio_path, title, artist, art_path)

    with open(audio_path, "rb") as f:
        await context.bot.send_audio(
            chat_id=chat_id,
            audio=f,
            title=title if title else audio_path.stem,
            performer=artist if artist else "فنان غير معروف",
            caption="✅ تم تحديث الـ Tags وحقن الغلاف بنجاح!",
        )

    for key in ["audio_state", "ready_audio_path", "meta_title", "meta_artist", "meta_art_path", "pending_audio", "pending_video"]:
        context.user_data.pop(key, None)


# ----------------------------------------------------------------------
# آليات معالجة تقسيم ودمج الـ PDF (مُنفَّذة بالكامل داخل executor الآن)
# ----------------------------------------------------------------------

async def split_pdf_pages(input_path: Path, output_path: Path, start_page: int, end_page: int) -> int:
    """يقص نطاق صفحات من PDF. يعيد عدد الصفحات الناتجة فعليًا (0 إن كان
    النطاق غير صالح)، ليتحقق المستدعي قبل إرسال ملف قد يكون فارغًا."""
    from pypdf import PdfReader, PdfWriter

    def _split_sync() -> int:
        reader = PdfReader(str(input_path))
        writer = PdfWriter()
        total_pages = len(reader.pages)
        start_idx = max(0, start_page - 1)
        end_idx = min(total_pages, end_page)
        for i in range(start_idx, end_idx):
            writer.add_page(reader.pages[i])
        page_count = max(0, end_idx - start_idx)
        if page_count > 0:
            with open(output_path, "wb") as f:
                writer.write(f)
        return page_count

    return await asyncio.get_running_loop().run_in_executor(None, _split_sync)


async def merge_pdf_files(input_paths: List[Path], output_path: Path):
    from pypdf import PdfWriter

    def _merge_sync():
        writer = PdfWriter()
        for path in input_paths:
            writer.append(str(path))
        with open(output_path, "wb") as f:
            writer.write(f)

    await asyncio.get_running_loop().run_in_executor(None, _merge_sync)


# ----------------------------------------------------------------------
# كيبورد اللوحات والقوائم التفاعلية
# ----------------------------------------------------------------------

def main_menu_keyboard() -> InlineKeyboardMarkup:
    buttons = [
        [InlineKeyboardButton("📂 أدوات تعديل الملفات", callback_data="sub_files")],
        [InlineKeyboardButton("🎵 أدوات تعديل الصوتيات", callback_data="sub_audio")],
    ]
    return InlineKeyboardMarkup(buttons)


def files_submenu_keyboard() -> InlineKeyboardMarkup:
    buttons = [
        [InlineKeyboardButton("📄 Word ➜ PDF", callback_data="mode_word2pdf")],
        [InlineKeyboardButton("📄 PDF ➜ Word", callback_data="mode_pdf2word")],
        [InlineKeyboardButton("✂️ قص صفحات PDF", callback_data="mode_split_pdf")],
        [InlineKeyboardButton("🔗 دمج ملفات PDF", callback_data="mode_merge_pdf")],
        [InlineKeyboardButton("📚 EPUB ➜ PDF", callback_data="mode_ebook")],
        [InlineKeyboardButton("🖼️ تحويل صور إلى PDF/Word", callback_data="mode_image")],
        [InlineKeyboardButton("🔒 تشفير حماية الـ PDF", callback_data="mode_encrypt_pdf")],
        [InlineKeyboardButton("🗜️ ضغط ملف PDF (تقليل الحجم)", callback_data="mode_compress_pdf")],
        [InlineKeyboardButton("🔙 العودة للقائمة الرئيسية", callback_data="back_to_main")],
    ]
    return InlineKeyboardMarkup(buttons)


def audio_submenu_keyboard() -> InlineKeyboardMarkup:
    buttons = [
        [InlineKeyboardButton("🎬 فيديو ➜ صوت MP3", callback_data="mode_video2audio")],
        [InlineKeyboardButton("🎵 تحويل صيغة صوتية", callback_data="mode_audio")],
        [InlineKeyboardButton("✂️ قص مقطع صوتي (Trim)", callback_data="mode_trim_audio")],
        [InlineKeyboardButton("⚡ تغيير سرعة الصوت", callback_data="mode_audio_speed")],
        [InlineKeyboardButton("🔗 دمج ملفات صوتية", callback_data="mode_merge_audio")],
        [InlineKeyboardButton("🔊 رفع / خفض الصوت", callback_data="mode_audio_volume")],
        [InlineKeyboardButton("🗣️ تحويل نص إلى صوت (TTS)", callback_data="mode_tts")],
        [InlineKeyboardButton("🔙 العودة للقائمة الرئيسية", callback_data="back_to_main")],
    ]
    return InlineKeyboardMarkup(buttons)


# ----------------------------------------------------------------------
# لوحة تحكم وإدارة الآدمن (Inline Admin Panel)
# ----------------------------------------------------------------------

def is_admin(user_id: int) -> bool:
    return user_id in ADMIN_IDS


def admin_keyboard() -> InlineKeyboardMarkup:
    buttons = [
        [InlineKeyboardButton("📊 إحصائيات البوت", callback_data="admin_stats")],
        [InlineKeyboardButton("📢 إذاعة جماعية (Broadcast)", callback_data="admin_broadcast")],
        [InlineKeyboardButton("🚫 حظر مستخدم", callback_data="admin_ban"),
         InlineKeyboardButton("🟢 إلغاء حظر", callback_data="admin_unban")],
        [InlineKeyboardButton("🔐 تعيين قناة الاشتراك الإجباري", callback_data="admin_set_sub")],
        [InlineKeyboardButton("🧹 مسح الكاش", callback_data="admin_clear_cache")],
    ]
    return InlineKeyboardMarkup(buttons)


async def admin_panel_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    await update.message.reply_text(
        "⚙️ <b>لوحة تحكم الإدارة وقاعدة البيانات</b>\nاختر من الأزرار الإجراء المطلوب:",
        reply_markup=admin_keyboard(),
        parse_mode=ParseMode.HTML,
    )


async def cancel_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """أمر /cancel لإلغاء أي عملية معلّقة (إدارية أو تحويل) بشكل آمن."""
    for key in (
        "admin_state", "pending_broadcast", "pdf_state", "audio_state", "current_mode",
        "merge_files", "merge_audio_files",
    ):
        context.user_data.pop(key, None)
    await update.message.reply_text("✅ تم إلغاء أي عملية معلّقة. أرسل /start للبدء من جديد.")


# ----------------------------------------------------------------------
# الإذاعة الجماعية (Broadcast) - إرسال متزامن محكوم مع معالجة الأخطاء
# ----------------------------------------------------------------------

async def _run_broadcast(context: ContextTypes.DEFAULT_TYPE, from_chat_id: int, message_id: int, status_message) -> None:
    pool = _get_pool()
    if pool is None:
        await status_message.edit_text("❌ قاعدة البيانات غير متاحة، تعذّر إتمام الإذاعة.")
        return

    try:
        async with pool.acquire() as conn:
            rows = await conn.fetch(
                "SELECT user_id FROM users WHERE is_banned = FALSE AND blocked_bot = FALSE;"
            )
    except Exception:
        logger.exception("فشل جلب قائمة المستخدمين للإذاعة")
        await status_message.edit_text("❌ فشل جلب قائمة المستخدمين.")
        return

    targets = [r["user_id"] for r in rows]
    total = len(targets)
    if total == 0:
        await status_message.edit_text("⚠️ لا يوجد مستخدمون لإرسال الإذاعة إليهم حاليًا.")
        return

    stats = {"success": 0, "blocked": 0, "failed": 0}
    semaphore = asyncio.Semaphore(BROADCAST_CONCURRENCY)

    async def _worker(target_id: int) -> None:
        async with semaphore:
            try:
                await context.bot.copy_message(chat_id=target_id, from_chat_id=from_chat_id, message_id=message_id)
                stats["success"] += 1
            except Forbidden:
                stats["blocked"] += 1
                try:
                    async with pool.acquire() as conn:
                        await conn.execute("UPDATE users SET blocked_bot = TRUE WHERE user_id = $1;", target_id)
                except Exception:
                    logger.exception(f"فشل تحديث حالة blocked_bot للمستخدم {target_id}")
            except RetryAfter as e:
                await asyncio.sleep(e.retry_after + 0.5)
                try:
                    await context.bot.copy_message(chat_id=target_id, from_chat_id=from_chat_id, message_id=message_id)
                    stats["success"] += 1
                except Exception:
                    stats["failed"] += 1
            except (BadRequest, TelegramError):
                stats["failed"] += 1
            except Exception:
                logger.exception(f"خطأ غير متوقع أثناء الإرسال إلى {target_id}")
                stats["failed"] += 1

            await asyncio.sleep(BROADCAST_DELAY_PER_MESSAGE)

    async def _progress_updater() -> None:
        try:
            while True:
                await asyncio.sleep(4)
                done = stats["success"] + stats["blocked"] + stats["failed"]
                with contextlib.suppress(Exception):
                    await status_message.edit_text(
                        f"⏳ جارٍ الإرسال... {done}/{total}\n"
                        f"✅ نجح: {stats['success']} | ⛔ محظور: {stats['blocked']} | ❌ فشل: {stats['failed']}"
                    )
        except asyncio.CancelledError:
            pass

    progress_task = asyncio.create_task(_progress_updater())
    try:
        await asyncio.gather(*(_worker(t) for t in targets))
    finally:
        progress_task.cancel()
        with contextlib.suppress(asyncio.CancelledError):
            await progress_task

    await status_message.edit_text(
        "📢 <b>اكتملت الإذاعة الجماعية</b>\n\n"
        f"✅ تم الإرسال بنجاح: <code>{stats['success']}</code>\n"
        f"⛔ حظروا البوت مسبقًا: <code>{stats['blocked']}</code>\n"
        f"❌ فشل لأسباب أخرى: <code>{stats['failed']}</code>",
        parse_mode=ParseMode.HTML,
    )
    await log_action("broadcast")


# ----------------------------------------------------------------------
# معالجة الرسائل والـ Callbacks الشاملة
# ----------------------------------------------------------------------

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    await register_user(user.id, user.username)

    if await _reject_if_banned(update):
        return

    missing_channel = await _is_subscription_required_and_unmet(user.id, context)
    if missing_channel:
        await update.message.reply_text(
            "⚠️ <b>عذرًا، يجب عليك الاشتراك في قناة البوت الرسمية أولاً لاستخدام ميزات "
            "معالجة وتحويل الملفات والصوتيات!</b>",
            reply_markup=get_sub_keyboard(missing_channel),
            parse_mode=ParseMode.HTML,
        )
        return

    context.user_data.clear()
    await update.message.reply_text(
        "👋 أهلًا بك في بوت تحويل الصيغ والوسائط المحترف!\n💡 اختر القسم المطلوب للبدء:",
        reply_markup=main_menu_keyboard(),
    )


async def menu_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    user_id = query.from_user.id

    if await is_user_banned(user_id):
        await query.answer("🚫 حسابك محظور من استخدام هذا البوت.", show_alert=True)
        return

    if query.data == "check_sub_again":
        missing_channel = await _is_subscription_required_and_unmet(user_id, context)
        if missing_channel:
            await query.answer("❌ يبدو أنك لم تشترك في القناة بعد! الرجاء الاشتراك والمحاولة مجددًا.", show_alert=True)
            return
        await query.answer("✅ شكراً لك على الاشتراك! تم فتح ميزات البوت الآن.", show_alert=True)
        context.user_data.clear()
        await query.edit_message_text("👋 أهلاً بك مجدداً! اختر القسم المطلوب للبدء الفوري:", reply_markup=main_menu_keyboard())
        return

    missing_channel = await _is_subscription_required_and_unmet(user_id, context)
    if missing_channel:
        await query.answer("⚠️ يجب عليك الاشتراك في القناة الرسمية أولاً!", show_alert=True)
        return

    await query.answer()
    data = query.data

    if data == "sub_files":
        await query.edit_message_text("📂 <b>قسم أدوات تعديل الملفات:</b>\nاختر الأداة المطلوبة:", reply_markup=files_submenu_keyboard(), parse_mode=ParseMode.HTML)
        return
    elif data == "sub_audio":
        await query.edit_message_text("🎵 <b>قسم أدوات تعديل الصوتيات:</b>\nاختر الأداة المطلوبة:", reply_markup=audio_submenu_keyboard(), parse_mode=ParseMode.HTML)
        return
    elif data == "back_to_main":
        await query.edit_message_text("👋 اختر القسم المطلوب من الأزرار أدناه للبدء:", reply_markup=main_menu_keyboard())
        return

    # معالجة أزرار الآدمن
    if data.startswith("admin_"):
        if not is_admin(user_id):
            await query.answer("❌ غير مسموح لك بالوصول.", show_alert=True)
            return
        await _handle_admin_callback(update, context, data)
        return

    # تفعيل المودات المختارة وتوجيه المستخدم
    context.user_data["current_mode"] = data
    if data == "mode_word2pdf":
        await query.edit_message_text("📄 أرسل الآن ملف Word (.doc أو .docx) لتحويله إلى PDF.")
    elif data == "mode_pdf2word":
        await query.edit_message_text("📄 أرسل الآن ملف PDF لتحويله إلى Word.")
    elif data == "mode_video2audio":
        await query.edit_message_text("🎬 أرسل ملف الفيديو لاستخراج الصوت منه بصيغة MP3.")
    elif data == "mode_ebook":
        await query.edit_message_text("📚 أرسل ملف EPUB ليتم تحويله تلقائيًا إلى PDF.")
    elif data == "mode_audio":
        await query.edit_message_text("🎵 أرسل الملف الصوتي المراد تعديله أو تحويل صيغته.")
    elif data == "mode_image":
        await query.edit_message_text("🖼️ أرسل الصورة أو مجموعة الصور المراد تجميعها.")
    elif data == "mode_encrypt_pdf":
        await query.edit_message_text("🔒 أرسل الآن ملف PDF لحمايته وتشفيره بكلمة مرور.")
    elif data == "mode_compress_pdf":
        await query.edit_message_text("🗜️ أرسل ملف الـ PDF الذي تود ضغطه وتقليص حجمه الآن.")
    elif data == "mode_split_pdf":
        await query.edit_message_text("✂️ أرسل أولاً ملف الـ PDF الذي ترغب بقص صفحات منه.")
    elif data == "mode_merge_pdf":
        context.user_data["pdf_state"] = "WAITING_MERGE_FILES"
        context.user_data["merge_files"] = []
        await query.edit_message_text(
            f"🔗 أرسل ملفات الـ PDF التي تود دمجها (حتى {MAX_MERGE_FILES} ملفًا)، "
            "وعند الانتهاء أرسل كلمة <b>دمج</b>.",
            parse_mode=ParseMode.HTML,
        )
    elif data == "mode_trim_audio":
        await query.edit_message_text("✂️ أرسل أولاً الملف الصوتي المراد قصه.")
    elif data == "mode_audio_speed":
        await query.edit_message_text("⚡ أرسل الملف الصوتي لتعديل وتغيير سرعته.")
    elif data == "mode_merge_audio":
        context.user_data["audio_state"] = "WAITING_MERGE_AUDIO"
        context.user_data["merge_audio_files"] = []
        await query.edit_message_text(
            f"🔗 أرسل الملفات الصوتية (حتى {MAX_MERGE_FILES} ملفًا)، ثم أرسل كلمة <b>دمج</b>.",
            parse_mode=ParseMode.HTML,
        )
    elif data == "mode_audio_volume":
        await query.edit_message_text("🔊 أرسل المقطع الصوتي المراد رفع أو خفض حجم ديسيبل الصوت له.")
    elif data == "mode_tts":
        context.user_data["audio_state"] = "WAITING_TTS_TEXT"
        await query.edit_message_text("🗣️ أرسل الآن النص (بالعربية أو الإنجليزية) لتحويله إلى مقطع صوتي مسموع.")


async def _handle_admin_callback(update: Update, context: ContextTypes.DEFAULT_TYPE, data: str) -> None:
    query = update.callback_query

    if data == "admin_stats":
        await _render_admin_stats(query)

    elif data == "admin_broadcast":
        context.user_data["admin_state"] = "WAITING_BROADCAST_MSG"
        await query.edit_message_text("📢 أرسل الآن رسالة الإذاعة (أو /cancel للإلغاء):")

    elif data == "admin_broadcast_confirm":
        pending = context.user_data.pop("pending_broadcast", None)
        if not pending:
            await query.edit_message_text("⚠️ لا توجد إذاعة معلّقة. ابدأ من جديد عبر /admin.")
            return
        await query.edit_message_text("⏳ جاري بدء الإذاعة الجماعية...")
        await _run_broadcast(context, pending["from_chat_id"], pending["message_id"], status_message=query.message)

    elif data == "admin_broadcast_cancel":
        context.user_data.pop("pending_broadcast", None)
        await query.edit_message_text("❌ تم إلغاء الإذاعة الجماعية.", reply_markup=admin_keyboard())

    elif data == "admin_ban":
        context.user_data["admin_state"] = "WAITING_BAN_ID"
        await query.edit_message_text("🚫 أرسل الـ User ID لحظره (أو /cancel للإلغاء):")

    elif data == "admin_unban":
        context.user_data["admin_state"] = "WAITING_UNBAN_ID"
        await query.edit_message_text("🟢 أرسل الـ User ID لإلغاء حظره (أو /cancel للإلغاء):")

    elif data == "admin_set_sub":
        context.user_data["admin_state"] = "WAITING_CHANNEL_USER"
        await query.edit_message_text(
            "🔐 <b>إعداد الاشتراك الإجباري:</b>\nأرسل الآن يوزر القناة الجديد (مثال: <code>@MyChannel</code>) "
            "أو اكتب <code>تعطيل</code> لإيقاف الميزة (أو /cancel للإلغاء):",
            parse_mode=ParseMode.HTML,
        )

    elif data == "admin_clear_cache":
        _ban_cache.invalidate()
        _settings_cache.invalidate()
        await query.edit_message_text("✅ تم مسح الكاش بنجاح.", reply_markup=admin_keyboard())


async def _render_admin_stats(query) -> None:
    await query.edit_message_text("⏳ جاري جلب الإحصائيات...")
    pool = _get_pool()
    if pool is None:
        await query.edit_message_text("❌ قاعدة البيانات غير متاحة حاليًا.", reply_markup=admin_keyboard())
        return
    try:
        async with pool.acquire() as conn:
            total_users = await conn.fetchval("SELECT COUNT(*) FROM users;")
            banned_users = await conn.fetchval("SELECT COUNT(*) FROM users WHERE is_banned = TRUE;")
            blocked_users = await conn.fetchval("SELECT COUNT(*) FROM users WHERE blocked_bot = TRUE;")
            active_24h = await conn.fetchval("SELECT COUNT(*) FROM users WHERE last_active > NOW() - INTERVAL '24 hours';")
            total_actions = await conn.fetchval("SELECT COUNT(*) FROM stats_log;")
        current_chan = await get_setting("force_channel") or "غير محددة ❌"
        stats_text = (
            "📊 <b>الإحصائيات:</b>\n"
            f"👥 مستخدمين: <code>{total_users}</code>\n"
            f"🟢 نشطون آخر 24 ساعة: <code>{active_24h}</code>\n"
            f"🚫 محظورين يدويًا: <code>{banned_users}</code>\n"
            f"⛔ حظروا البوت: <code>{blocked_users}</code>\n"
            f"⚙️ عمليات ناجحة: <code>{total_actions}</code>\n"
            f"📢 القناة الحالية: <code>{html.escape(current_chan)}</code>"
        )
        await query.edit_message_text(stats_text, reply_markup=admin_keyboard(), parse_mode=ParseMode.HTML)
    except Exception:
        logger.exception("خطأ أثناء جلب الإحصائيات")
        await query.edit_message_text("❌ حدث خطأ أثناء جلب الإحصائيات.", reply_markup=admin_keyboard())


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await _reject_if_banned(update):
        return

    user_id = update.effective_user.id
    missing_channel = await _is_subscription_required_and_unmet(user_id, context)
    if missing_channel:
        await update.message.reply_text(
            "⚠️ لا يمكنك إرسال ملفات! يجب الاشتراك بقناة البوت الرسمية أولاً.",
            reply_markup=get_sub_keyboard(missing_channel),
        )
        return

    if await _reject_if_rate_limited(update):
        return

    doc = update.message.document
    if not await _check_size_or_reject(update, doc.file_size):
        return

    filename = doc.file_name or "file"
    mode = context.user_data.get("current_mode")

    # معالجة ضغط PDF
    if mode == "mode_compress_pdf" and filename.lower().endswith(".pdf"):
        msg = await update.message.reply_text("⏳ جاري تهيئة وتحميل ملف الـ PDF لبدء الضغط...")
        try:
            lp = await download_telegram_file(context, doc.file_id, doc.file_unique_id, filename)
            out_p = CONVERTED_DIR / f"compressed_{filename}"
            async with HeavyTaskGuard():
                await compress_pdf_file_async(lp, out_p, msg)
            await log_action("compress_pdf")
            with open(out_p, "rb") as f:
                await update.message.reply_document(document=f, filename=out_p.name, caption="✅ تم ضغط الملف وتقليل الحجم بنجاح!")
        except Exception as e:
            logger.exception("فشل ضغط ملف PDF")
            await update.message.reply_text("❌ تعذر ضغط هذا الملف بسبب قيود في بنيته الداخلية.")
        finally:
            with contextlib.suppress(Exception):
                await msg.delete()
        return

    # معالجة استلام ملف PDF تمهيدًا لقصّه
    if mode == "mode_split_pdf" and filename.lower().endswith(".pdf"):
        lp = await download_telegram_file(context, doc.file_id, doc.file_unique_id, filename)
        context.user_data["split_pdf_source"] = str(lp)
        context.user_data["pdf_state"] = "WAITING_SPLIT_RANGE"
        await update.message.reply_text(
            "⏱️ تم حفظ الملف. أرسل الآن نطاق الصفحات المطلوب قصها هكذا:\n<code>1-15</code>",
            parse_mode=ParseMode.HTML,
        )
        return

    # معالجة استقبال ملفات PDF المتعددة للدمج
    if mode == "mode_merge_pdf" and filename.lower().endswith(".pdf"):
        merge_list = context.user_data.setdefault("merge_files", [])
        if len(merge_list) >= MAX_MERGE_FILES:
            await update.message.reply_text(
                f"⚠️ وصلت للحد الأقصى ({MAX_MERGE_FILES} ملفًا). أرسل كلمة <b>دمج</b> لإتمام العملية الآن.",
                parse_mode=ParseMode.HTML,
            )
            return
        lp = await download_telegram_file(context, doc.file_id, doc.file_unique_id, filename)
        merge_list.append(str(lp))
        await update.message.reply_text(
            f"📥 تم استقبال الملف رقم ({len(merge_list)}).\nأرسل التالي، أو أرسل كلمة <b>دمج</b> لإتمام العملية.",
            parse_mode=ParseMode.HTML,
        )
        return

    # معالجة استلام ملف صوتي كـ Document
    if Path(filename).suffix.lower() in AUDIO_EXTENSIONS:
        if mode == "mode_audio_speed":
            lp = await download_telegram_file(context, doc.file_id, doc.file_unique_id, filename)
            context.user_data["speed_source_path"] = str(lp)
            context.user_data["audio_state"] = "WAITING_SPEED_VALUE"
            await update.message.reply_text(
                "⏱️ أرسل سرعة المعالجة الصوتية كرقم عشري بين <code>0.5</code> و<code>2.0</code>:",
                parse_mode=ParseMode.HTML,
            )
            return
        elif mode == "mode_merge_audio":
            merge_list = context.user_data.setdefault("merge_audio_files", [])
            if len(merge_list) >= MAX_MERGE_FILES:
                await update.message.reply_text(
                    f"⚠️ وصلت للحد الأقصى ({MAX_MERGE_FILES} ملفًا). أرسل <b>دمج</b> الآن.", parse_mode=ParseMode.HTML
                )
                return
            lp = await download_telegram_file(context, doc.file_id, doc.file_unique_id, filename)
            merge_list.append(str(lp))
            await update.message.reply_text(f"📥 تم حفظ المقطع رقم ({len(merge_list)}). أرسل التالي أو اكتب <b>دمج</b>.", parse_mode=ParseMode.HTML)
            return
        elif mode == "mode_audio_volume":
            lp = await download_telegram_file(context, doc.file_id, doc.file_unique_id, filename)
            context.user_data["volume_source_path"] = str(lp)
            context.user_data["audio_state"] = "WAITING_VOLUME_VALUE"
            await update.message.reply_text("🔊 أرسل مستوى الصوت بالديسيبل (dB)، مثلاً 6 أو -6:")
            return
        elif mode == "mode_trim_audio":
            lp = await download_telegram_file(context, doc.file_id, doc.file_unique_id, filename)
            context.user_data["trim_source_path"] = str(lp)
            context.user_data["audio_state"] = "WAITING_TRIM_TIME"
            await update.message.reply_text(
                "⏱️ أرسل توقيت القص هكذا تمامًا:\n<code>00:01:10 - 00:02:45</code>", parse_mode=ParseMode.HTML
            )
            return

    is_handled = await files_handler.handle_files_document(update, context)
    if not is_handled:
        await audio_handler.handle_audio_document(update, context)


async def handle_audio_message_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await _reject_if_banned(update):
        return

    user_id = update.effective_user.id
    missing_channel = await _is_subscription_required_and_unmet(user_id, context)
    if missing_channel:
        await update.message.reply_text(
            "⚠️ لا يمكنك إرسال وسائط! يرجى الاشتراك أولاً.", reply_markup=get_sub_keyboard(missing_channel)
        )
        return

    if await _reject_if_rate_limited(update):
        return

    audio = update.message.audio or update.message.voice
    if not await _check_size_or_reject(update, getattr(audio, "file_size", None)):
        return

    mode = context.user_data.get("current_mode")
    filename = getattr(audio, "file_name", None) or "voice.ogg"

    if mode == "mode_trim_audio":
        lp = await download_telegram_file(context, audio.file_id, audio.file_unique_id, filename)
        context.user_data["trim_source_path"] = str(lp)
        context.user_data["audio_state"] = "WAITING_TRIM_TIME"
        await update.message.reply_text(
            "⏱️ أرسل توقيت القص هكذا تمامًا:\n<code>00:01:10 - 00:02:45</code>", parse_mode=ParseMode.HTML
        )
        return

    elif mode == "mode_audio_speed":
        lp = await download_telegram_file(context, audio.file_id, audio.file_unique_id, filename)
        context.user_data["speed_source_path"] = str(lp)
        context.user_data["audio_state"] = "WAITING_SPEED_VALUE"
        await update.message.reply_text("⏱️ أرسل سرعة المعالجة كرقم عشري (مثلاً 1.25):")
        return

    elif mode == "mode_merge_audio":
        merge_list = context.user_data.setdefault("merge_audio_files", [])
        if len(merge_list) >= MAX_MERGE_FILES:
            await update.message.reply_text(f"⚠️ وصلت للحد الأقصى ({MAX_MERGE_FILES} ملفًا). أرسل <b>دمج</b> الآن.", parse_mode=ParseMode.HTML)
            return
        lp = await download_telegram_file(context, audio.file_id, audio.file_unique_id, filename)
        merge_list.append(str(lp))
        await update.message.reply_text(f"📥 تم حفظ المقطع رقم ({len(merge_list)}). أرسل التالي أو اكتب <b>دمج</b>.", parse_mode=ParseMode.HTML)
        return

    elif mode == "mode_audio_volume":
        lp = await download_telegram_file(context, audio.file_id, audio.file_unique_id, filename)
        context.user_data["volume_source_path"] = str(lp)
        context.user_data["audio_state"] = "WAITING_VOLUME_VALUE"
        await update.message.reply_text("🔊 أرسل القيمة بالديسيبل (dB): مثلاً 5 أو -5:")
        return

    await audio_handler.handle_audio_message(update, context)


async def handle_video_message_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """موجّه مركزي لرسائل الفيديو، يطبّق نفس حمايات handle_audio_message_router
    (حظر، اشتراك إجباري، Rate Limit، حد الحجم) قبل التفويض لـ audio_handler.
    هذا الموجّه ضروري لأن استدعاء audio_handler.handle_video_message مباشرة
    (كما كان سابقًا) كان يتجاوز كل هذه الحمايات المركزية تمامًا."""
    if await _reject_if_banned(update):
        return

    user_id = update.effective_user.id
    missing_channel = await _is_subscription_required_and_unmet(user_id, context)
    if missing_channel:
        await update.message.reply_text(
            "⚠️ لا يمكنك إرسال وسائط! يرجى الاشتراك أولاً.", reply_markup=get_sub_keyboard(missing_channel)
        )
        return

    if await _reject_if_rate_limited(update):
        return

    video = update.message.video or update.message.video_note
    if not await _check_size_or_reject(update, getattr(video, "file_size", None)):
        return

    await audio_handler.handle_video_message(update, context)


async def handle_text_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await _reject_if_banned(update):
        return

    user_id = update.effective_user.id
    admin_state = context.user_data.get("admin_state")
    state = context.user_data.get("audio_state")
    pdf_state = context.user_data.get("pdf_state")
    text = (update.message.text or "").strip()

    # ---------------- إعداد قناة الاشتراك الإجباري ----------------
    if admin_state == "WAITING_CHANNEL_USER" and is_admin(user_id):
        context.user_data.pop("admin_state", None)
        if text in ("تعطيل", "disable"):
            await set_setting("force_channel", "")
            await update.message.reply_text("🟢 تم تعطيل نظام الاشتراك الإجباري بنجاح.")
        else:
            if not text.startswith("@"):
                text = f"@{text}"
            await set_setting("force_channel", text)
            await update.message.reply_text(
                f"🔐 تم تفعيل القناة بنجاح!\nيوزر القناة: <code>{html.escape(text)}</code>\n"
                "تأكد من رفع البوت كمشرف داخلها.",
                parse_mode=ParseMode.HTML,
            )
        return

    # ---------------- نطاق قص صفحات PDF ----------------
    if pdf_state == "WAITING_SPLIT_RANGE":
        if "-" not in text:
            await update.message.reply_text("⚠️ أرسل النطاق بصيغة صحيحة، مثال: 1-15")
            return
        src_path_str = context.user_data.get("split_pdf_source")
        if not src_path_str:
            await update.message.reply_text("❌ لم يتم العثور على ملف PDF المرفوع مسبقًا، أعد المحاولة.")
            return
        context.user_data.pop("pdf_state", None)
        try:
            start_p, end_p = map(int, text.split("-", maxsplit=1))
        except ValueError:
            await update.message.reply_text("⚠️ أدخل أرقامًا صحيحة، مثال: 1-15")
            return
        if start_p < 1 or end_p < start_p:
            await update.message.reply_text("⚠️ نطاق غير صالح. تأكد أن رقم الصفحة الأولى ≥ 1 وأصغر من أو يساوي الأخيرة.")
            return

        msg = await update.message.reply_text("⏳ جاري قص النطاق المحدد...")
        src_p = Path(src_path_str)
        out_p = CONVERTED_DIR / f"clipped_{start_p}_to_{end_p}_{src_p.name}"
        try:
            page_count = await split_pdf_pages(src_p, out_p, start_p, end_p)
            if page_count == 0:
                await msg.edit_text("⚠️ النطاق المُدخل خارج عدد صفحات الملف الفعلي.")
                return
            await log_action("pdf_split")
            with open(out_p, "rb") as f:
                await update.message.reply_document(document=f, filename=out_p.name, caption=f"✂️ تم قص الصفحات من {start_p} إلى {end_p} بنجاح!")
            await msg.delete()
        except Exception as e:
            logger.exception("خطأ أثناء قص PDF")
            await msg.edit_text(f"❌ حدث خطأ أثناء المعالجة: {html.escape(str(e))}")
        return

    # ---------------- دمج ملفات PDF ----------------
    if pdf_state == "WAITING_MERGE_FILES" and text in ("دمج", "merge"):
        files_list = context.user_data.get("merge_files", [])
        if len(files_list) < 2:
            await update.message.reply_text("⚠️ يجب إرسال ملفين PDF على الأقل لدمجهما.")
            return
        context.user_data.pop("pdf_state", None)
        msg = await update.message.reply_text(f"🔗 جاري دمج {len(files_list)} ملفات PDF...")
        out_p = CONVERTED_DIR / f"merged_document_{int(time.time())}.pdf"
        try:
            async with HeavyTaskGuard():
                await merge_pdf_files([Path(p) for p in files_list], out_p)
            await log_action("pdf_merge")
            with open(out_p, "rb") as f:
                await update.message.reply_document(document=f, filename="Merged_Document.pdf", caption="🔗 تم دمج جميع الملفات بنجاح!")
            await msg.delete()
            context.user_data.pop("merge_files", None)
        except Exception as e:
            logger.exception("خطأ أثناء دمج PDF")
            await msg.edit_text(f"❌ حدث خطأ أثناء الدمج: {html.escape(str(e))}")
        return

    # ---------------- دمج ملفات صوتية ----------------
    if state == "WAITING_MERGE_AUDIO" and text in ("دمج", "merge"):
        audio_list = context.user_data.get("merge_audio_files", [])
        if len(audio_list) < 2:
            await update.message.reply_text("⚠️ يجب إرسال ملفين صوتيين على الأقل لدمجهما.")
            return
        context.user_data.pop("audio_state", None)
        msg = await update.message.reply_text("⏳ جاري دمج المقاطع الصوتية...")
        out_p = CONVERTED_DIR / f"merged_audio_{int(time.time())}.mp3"
        try:
            async with HeavyTaskGuard():
                await merge_audio_files([Path(p) for p in audio_list], out_p)
            await log_action("audio_merge")
            with open(out_p, "rb") as f:
                await update.message.reply_audio(audio=f, caption="🔗 تم دمج كافة المقاطع بنجاح!")
            await msg.delete()
            context.user_data.pop("merge_audio_files", None)
        except Exception as e:
            logger.exception("خطأ أثناء دمج الصوت")
            await msg.edit_text(f"❌ حدث خطأ أثناء الدمج: {html.escape(str(e))}")
        return

    # ---------------- تغيير سرعة الصوت ----------------
    if state == "WAITING_SPEED_VALUE":
        try:
            speed_val = float(text)
            if not (0.5 <= speed_val <= 2.0):
                await update.message.reply_text("⚠️ القيمة يجب أن تكون بين 0.5 و2.0.")
                return
        except ValueError:
            await update.message.reply_text("⚠️ أدخل رقمًا عشريًا صحيحًا (مثال: 1.5)")
            return
        context.user_data.pop("audio_state", None)
        src_path = context.user_data.get("speed_source_path")
        if not src_path:
            await update.message.reply_text("❌ لم يتم العثور على الملف الأصلي.")
            return
        msg = await update.message.reply_text("⏳ جاري تعديل السرعة...")
        src_p = Path(src_path)
        out_p = CONVERTED_DIR / f"speed_{speed_val}_{src_p.name}"
        try:
            await change_audio_speed(src_p, out_p, speed_val)
            await log_action("audio_speed")
            with open(out_p, "rb") as f:
                await update.message.reply_audio(audio=f, caption=f"⚡ تم تعديل السرعة إلى {speed_val}x بنجاح!")
            await msg.delete()
        except Exception as e:
            logger.exception("خطأ أثناء تعديل سرعة الصوت")
            await msg.edit_text(f"❌ حدث خطأ: {html.escape(str(e))}")
        return

    # ---------------- تعديل حجم الصوت ----------------
    if state == "WAITING_VOLUME_VALUE":
        try:
            vol_val = float(text)
        except ValueError:
            await update.message.reply_text("⚠️ أدخل رقمًا صحيحًا (مثال: 5 أو -5)")
            return
        context.user_data.pop("audio_state", None)
        src_path = context.user_data.get("volume_source_path")
        if not src_path:
            await update.message.reply_text("❌ لم يتم العثور على الملف الأصلي.")
            return
        msg = await update.message.reply_text("⏳ جاري تعديل مستوى الصوت...")
        src_p = Path(src_path)
        out_p = CONVERTED_DIR / f"vol_{vol_val}_{src_p.name}"
        try:
            await change_audio_volume(src_p, out_p, vol_val)
            await log_action("audio_volume")
            with open(out_p, "rb") as f:
                await update.message.reply_audio(audio=f, caption=f"🔊 تم تعديل حجم الصوت بمقدار {vol_val}dB بنجاح!")
            await msg.delete()
        except Exception as e:
            logger.exception("خطأ أثناء تعديل حجم الصوت")
            await msg.edit_text(f"❌ حدث خطأ: {html.escape(str(e))}")
        return

    # ---------------- تحويل نص إلى صوت (TTS) ----------------
    if state == "WAITING_TTS_TEXT":
        if await _reject_if_rate_limited(update):
            return
        context.user_data.pop("audio_state", None)
        if len(text) > 2000:
            await update.message.reply_text("⚠️ النص طويل جدًا (الحد الأقصى 2000 حرف).")
            return
        msg = await update.message.reply_text("⏳ جاري توليد المقطع الصوتي...")
        try:
            from gtts import gTTS
            out_p = CONVERTED_DIR / f"tts_{update.message.message_id}.mp3"
            lang = "ar" if any("\u0600" <= c <= "\u06FF" for c in text) else "en"

            def _generate_tts():
                tts = gTTS(text=text, lang=lang, slow=False)
                tts.save(str(out_p))

            async with HeavyTaskGuard():
                # gTTS يعتمد على استدعاء شبكي خارجي (Google Translate)؛
                # نضبط مهلة قصوى لتفادي تعليق الطلب إلى الأبد إن تعطلت الخدمة.
                await asyncio.wait_for(
                    asyncio.get_running_loop().run_in_executor(None, _generate_tts),
                    timeout=60.0,
                )
            await log_action("tts_generated")
            with open(out_p, "rb") as f:
                await update.message.reply_audio(audio=f, caption="🗣️ تم توليد النطق الصوتي بنجاح!")
            await msg.delete()
        except Exception as e:
            logger.exception("خطأ أثناء توليد TTS")
            await msg.edit_text(f"❌ فشل توليد الصوت: {html.escape(str(e))}")
        return

    # ---------------- قص مقطع صوتي ----------------
    if state == "WAITING_TRIM_TIME":
        if " - " not in text:
            await update.message.reply_text("⚠️ أرسل التوقيت هكذا تمامًا:\n00:00:10 - 00:00:40")
            return
        context.user_data.pop("audio_state", None)
        start_t, end_t = (p.strip() for p in text.split(" - ", maxsplit=1))
        src_path = context.user_data.get("trim_source_path")
        if not src_path:
            await update.message.reply_text("❌ لم يتم العثور على الملف الصوتي الأصلي.")
            return
        msg = await update.message.reply_text("⏳ جاري قص المقطع...")
        src_p = Path(src_path)
        out_p = CONVERTED_DIR / f"trimmed_{src_p.name}"
        try:
            await trim_audio_file(src_p, out_p, start_t, end_t)
            await log_action("audio_trim")
            with open(out_p, "rb") as f:
                await update.message.reply_audio(audio=f, caption=f"✂️ تم قص المقطع من {start_t} إلى {end_t}!")
            await msg.delete()
        except Exception as e:
            logger.exception("خطأ أثناء قص الصوت")
            await msg.edit_text(f"❌ خطأ: {html.escape(str(e))}")
        return

    # ---------------- أوامر الأدمن (بث، حظر، إلغاء حظر) ----------------
    if admin_state and is_admin(user_id):
        if admin_state == "WAITING_BROADCAST_MSG":
            context.user_data.pop("admin_state", None)
            context.user_data["pending_broadcast"] = {
                "from_chat_id": update.message.chat_id,
                "message_id": update.message.message_id,
            }
            pool = _get_pool()
            recipient_count = 0
            if pool is not None:
                try:
                    async with pool.acquire() as conn:
                        recipient_count = await conn.fetchval(
                            "SELECT COUNT(*) FROM users WHERE is_banned = FALSE AND blocked_bot = FALSE;"
                        )
                except Exception:
                    logger.exception("فشل حساب عدد مستلمي الإذاعة")

            confirm_buttons = InlineKeyboardMarkup([[
                InlineKeyboardButton("✅ تأكيد الإرسال", callback_data="admin_broadcast_confirm"),
                InlineKeyboardButton("❌ إلغاء", callback_data="admin_broadcast_cancel"),
            ]])
            await update.message.reply_text(
                f"📢 سيتم الإرسال إلى <b>{recipient_count}</b> مستخدم. هل تريد المتابعة؟",
                reply_markup=confirm_buttons,
                parse_mode=ParseMode.HTML,
            )
            return

        elif admin_state == "WAITING_BAN_ID":
            context.user_data.pop("admin_state", None)
            try:
                target_id = int(text)
            except ValueError:
                await update.message.reply_text("⚠️ المعرف يجب أن يكون رقمًا صحيحًا.")
                return
            try:
                await set_user_ban_status(target_id, True)
                await update.message.reply_text(f"✅ تم حظر المستخدم <code>{target_id}</code>.", parse_mode=ParseMode.HTML)
            except Exception:
                logger.exception(f"فشل حظر المستخدم {target_id}")
                await update.message.reply_text("❌ حدث خطأ أثناء الحظر.")
            return

        elif admin_state == "WAITING_UNBAN_ID":
            context.user_data.pop("admin_state", None)
            try:
                target_id = int(text)
            except ValueError:
                await update.message.reply_text("⚠️ المعرف يجب أن يكون رقمًا صحيحًا.")
                return
            try:
                await set_user_ban_status(target_id, False)
                await update.message.reply_text(f"🟢 تم إلغاء حظر المستخدم <code>{target_id}</code>.", parse_mode=ParseMode.HTML)
            except Exception:
                logger.exception(f"فشل إلغاء حظر المستخدم {target_id}")
                await update.message.reply_text("❌ حدث خطأ أثناء إلغاء الحظر.")
            return

    # ---------------- تسليم إلى files_handler لحالة كلمة مرور PDF ----------------
    if pdf_state == "WAITING_PASSWORD":
        await files_handler.process_pdf_encryption(update, context, text)
        return

    # ---------------- تعديل الميتاداتا الصوتية ----------------
    if state == "WATING_TITLE":
        context.user_data["meta_title"] = text
        context.user_data["audio_state"] = "WATING_ARTIST"
        await update.message.reply_text("👤 أرسل الآن اسم الفنان:", reply_markup=audio_handler.metadata_skip_keyboard())
    elif state == "WATING_ARTIST":
        context.user_data["meta_artist"] = text
        context.user_data["audio_state"] = "WATING_ART"
        await update.message.reply_text("🖼️ أرسل الآن صورة الغلاف:", reply_markup=audio_handler.metadata_skip_keyboard())


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await _reject_if_banned(update):
        return

    user_id = update.effective_user.id
    missing_channel = await _is_subscription_required_and_unmet(user_id, context)
    if missing_channel:
        await update.message.reply_text("⚠️ لا يمكنك إرسال صور! يرجى الاشتراك أولاً.", reply_markup=get_sub_keyboard(missing_channel))
        return

    if await _reject_if_rate_limited(update):
        return

    photo = update.message.photo[-1]
    if not await _check_size_or_reject(update, getattr(photo, "file_size", None)):
        return

    if context.user_data.get("audio_state") == "WATING_ART":
        await audio_handler.handle_photo_as_art(update, context, photo)
        return
    await files_handler.handle_files_photo(update, context)


async def cleanup_job(context: ContextTypes.DEFAULT_TYPE):
    now = time.time()
    removed = 0
    for folder in (DOWNLOADS_DIR, CONVERTED_DIR):
        for path in folder.glob("*"):
            try:
                if path.is_file() and (now - path.stat().st_mtime) > FILE_MAX_AGE:
                    path.unlink()
                    removed += 1
            except OSError:
                pass
    if removed:
        logger.info(f"🧹 تم تنظيف {removed} ملف مؤقت.")

    pruned = prune_rate_limiter()
    if pruned:
        logger.info(f"🧹 تم تنظيف {pruned} سجل غير نشط من محدد المعدل.")


async def global_error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    """معالج أخطاء عام: يمنع صمت البوت التام عند حدوث استثناء غير متوقع."""
    logger.error("حدث استثناء أثناء معالجة تحديث:", exc_info=context.error)
    if isinstance(update, Update) and update.effective_message:
        with contextlib.suppress(Exception):
            await update.effective_message.reply_text(
                "❌ حدث خطأ غير متوقع أثناء معالجة طلبك. الرجاء المحاولة مجددًا لاحقًا."
            )


# ----------------------------------------------------------------------
# دالة بدء التشغيل الأساسية للبوت وإطلاق قاعدة البيانات
# ----------------------------------------------------------------------

async def _post_init(application: Application) -> None:
    success = await init_db()
    if not success:
        logger.critical("⚠️ البوت سيعمل بدون قاعدة بيانات فعالة؛ ميزات الأدمن والاشتراك الإجباري ستكون معطّلة.")


async def _post_shutdown(application: Application) -> None:
    await close_db()


def main():
    request_config = HTTPXRequest(
        connect_timeout=30.0,
        read_timeout=60.0,
        write_timeout=30.0,
        # Pool أكبر يسمح بعدد أعلى من الطلبات المتزامنة (مهم أثناء الإذاعة الجماعية
        # أو استخدام عدة مستخدمين للبوت في نفس اللحظة)
        connection_pool_size=100,
    )

    app = (
        Application.builder()
        .token(BOT_TOKEN)
        .request(request_config)
        .post_init(_post_init)
        .post_shutdown(_post_shutdown)
        .build()
    )

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("admin", admin_panel_command))
    app.add_handler(CommandHandler("cancel", cancel_command))

    # ⚠️ ترتيب مهم جدًا: يجب تسجيل معالِجات الأزرار ذات الأنماط المحددة
    # (audiofmt_, vidtarget_extract, meta_skip, وما يعادلها في files_handler)
    # قبل menu_callback العام الذي يستقبل CallbackQueryHandler بلا أي نمط
    # (pattern=None) ويطابق تاليًا أي بيانات زر. مكتبة python-telegram-bot
    # تُجرّب المعالِجات ضمن نفس المجموعة بترتيب التسجيل وتتوقف عند أول
    # تطابق؛ فلو سُجّل menu_callback أولًا فسيلتهم كل الضغطات (بما فيها
    # "meta_skip" الخاص بخطوات العنوان/الفنان/الغلاف) قبل أن تصل لمعالِجها
    # الصحيح في audio_handler.py، مما كان يُبقي حالة "audio_state" عالقة
    # ولا تتقدم أبدًا إلى "WATING_ART"، فتُعامَل صورة الغلاف كصورة عادية
    # بدل تطبيقها كغلاف فعلي على الملف الصوتي.
    files_handler.register_files_handlers(app)
    audio_handler.register_audio_handlers(app)

    app.add_handler(CallbackQueryHandler(menu_callback))

    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text_input))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.AUDIO | filters.VOICE, handle_audio_message_router))
    app.add_handler(MessageHandler(filters.VIDEO | filters.VIDEO_NOTE, handle_video_message_router))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))

    app.add_error_handler(global_error_handler)

    app.job_queue.run_repeating(cleanup_job, interval=CLEANUP_INTERVAL, first=CLEANUP_INTERVAL)

    logger.info("🚀 البوت انطلق رسميًا بكافة ميزات الضغط، الاشتراك الإجباري والتحكم الصوتي المتقدم...")
    app.run_polling(allowed_updates=Update.ALL_TYPES, drop_pending_updates=True)


if __name__ == "__main__":
    main()
